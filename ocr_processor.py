#!/usr/bin/env python3
"""
Enhanced OCR Processor with:
- Smart document naming
- Backup system
- Multi-page detection
- Text formatting preservation
- OCR confidence scoring
"""

import os
import shutil
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

# Configuration
OCR_FOLDER = "/mnt/f/OneDrive - UAB/Imágenes/Álbum de la cámara/AI_OCR"
OUTPUT_FOLDER = "/mnt/f/OneDrive - UAB/AI_Assistant"
PROCESSED_FOLDER = os.path.join(OCR_FOLDER, "Processed")
LOG_FILE = "/home/tayyabcheema777/ali/ocr_processing.log"
PROCESSED_LOG = "/home/tayyabcheema777/ali/processed_images.log"

def log_message(message):
    """Log messages with timestamp"""
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    with open(LOG_FILE, 'a') as f:
        f.write(f"[{timestamp}] {message}\n")
    print(f"[{timestamp}] {message}")

def get_processed_images():
    """Get list of already processed images"""
    if os.path.exists(PROCESSED_LOG):
        with open(PROCESSED_LOG, 'r') as f:
            return set(line.strip() for line in f)
    return set()

def mark_as_processed(filename):
    """Mark image as processed"""
    with open(PROCESSED_LOG, 'a') as f:
        f.write(f"{filename}\n")

def extract_first_words(text, word_count=3):
    """Extract first few words for filename"""
    words = re.findall(r'\b\w+\b', text)
    first_words = '_'.join(words[:word_count])
    # Clean for filename
    first_words = re.sub(r'[^\w\s-]', '', first_words)
    return first_words[:30]  # Limit length

def detect_list_formatting(text):
    """Detect if text contains numbered lists or bullet points"""
    has_numbered = bool(re.search(r'^\s*\d+[\)\.:]', text, re.MULTILINE))
    has_bullets = bool(re.search(r'^\s*[-•*]', text, re.MULTILINE))
    return has_numbered, has_bullets

def preserve_formatting(doc, text):
    """Add text to document with preserved formatting"""
    lines = text.split('\n')

    for line in lines:
        if not line.strip():
            doc.add_paragraph()
            continue

        # Check if numbered list
        numbered_match = re.match(r'^(\s*)(\d+[\)\.:])\s*(.+)$', line)
        if numbered_match:
            indent, marker, content = numbered_match.groups()
            p = doc.add_paragraph(content, style='List Number')
            continue

        # Check if bullet point
        bullet_match = re.match(r'^(\s*)[-•*]\s*(.+)$', line)
        if bullet_match:
            indent, content = bullet_match.groups()
            p = doc.add_paragraph(content, style='List Bullet')
            continue

        # Regular paragraph with indentation detection
        indent_match = re.match(r'^(\s+)(.+)$', line)
        if indent_match:
            spaces, content = indent_match.groups()
            p = doc.add_paragraph(content)
            # Add indentation (roughly 0.5 inch per 4 spaces)
            indent_level = len(spaces) // 4
            p.paragraph_format.left_indent = Inches(0.5 * indent_level)
        else:
            doc.add_paragraph(line)

def calculate_confidence(text):
    """Simple confidence calculation based on text characteristics"""
    if not text or len(text) < 10:
        return 0.3, "Very short text - low confidence"

    # Check for common indicators
    word_count = len(text.split())
    has_proper_punctuation = bool(re.search(r'[.,!?;:]', text))
    has_proper_capitalization = bool(re.search(r'[A-Z]', text))

    confidence = 0.5
    reasons = []

    if word_count > 20:
        confidence += 0.2
    else:
        reasons.append("Short text")

    if has_proper_punctuation:
        confidence += 0.15
    else:
        reasons.append("Limited punctuation")

    if has_proper_capitalization:
        confidence += 0.15
    else:
        reasons.append("Limited capitalization")

    reason_text = ", ".join(reasons) if reasons else "Good text characteristics"
    return min(confidence, 1.0), reason_text

def process_images():
    """Main processing function"""
    # Get today's date
    today = datetime.now().strftime("%Y%m%d")

    # Find new images
    processed_images = get_processed_images()
    new_images = []

    for filename in os.listdir(OCR_FOLDER):
        if filename.startswith(f"WIN_{today}") and filename.endswith(".jpg"):
            if filename not in processed_images:
                new_images.append(filename)

    if not new_images:
        log_message(f"No new images found for {today}")
        return

    new_images.sort()  # Process in order
    log_message(f"Found {len(new_images)} new image(s) to process")

    # Create document
    doc = Document()

    # Add title
    title = doc.add_paragraph(f"Transcribed Notes - {datetime.now().strftime('%B %d, %Y')}")
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(16)
    doc.add_paragraph()

    all_text = []
    total_confidence = 0

    for idx, image_file in enumerate(new_images, 1):
        image_path = os.path.join(OCR_FOLDER, image_file)

        log_message(f"Processing image {idx}/{len(new_images)}: {image_file}")

        # Here Claude would extract text using Read tool
        # For now, placeholder - in actual workflow, this would be replaced with actual OCR
        extracted_text = f"[Text from {image_file} would be extracted here by Claude]"

        # Apply British English corrections and punctuation
        # (This would be done by Claude's vision capabilities)

        # Calculate confidence
        confidence, reason = calculate_confidence(extracted_text)
        total_confidence += confidence

        # Add page header if multiple images
        if len(new_images) > 1:
            page_header = doc.add_paragraph(f"Page {idx} of {len(new_images)}")
            page_header.runs[0].bold = True
            page_header.runs[0].font.size = Pt(12)
            doc.add_paragraph()

        # Add text with formatting
        preserve_formatting(doc, extracted_text)

        # Add confidence note if low
        if confidence < 0.7:
            confidence_note = doc.add_paragraph(f"[OCR Confidence: {confidence:.0%} - {reason}]")
            confidence_note.runs[0].italic = True
            confidence_note.runs[0].font.size = Pt(9)

        doc.add_paragraph()  # Separator between pages
        all_text.append(extracted_text)

        # Mark as processed and backup
        mark_as_processed(image_file)

        # Move to processed folder
        os.makedirs(PROCESSED_FOLDER, exist_ok=True)
        backup_path = os.path.join(PROCESSED_FOLDER, image_file)
        shutil.move(image_path, backup_path)
        log_message(f"Moved {image_file} to Processed folder")

    # Generate smart filename
    combined_text = " ".join(all_text)
    first_words = extract_first_words(combined_text)
    page_suffix = f"_{len(new_images)}pages" if len(new_images) > 1 else ""
    filename = f"transcribed_{today}_{first_words}{page_suffix}.docx"

    # Add summary at the end
    doc.add_paragraph()
    summary = doc.add_paragraph("─" * 50)
    summary = doc.add_paragraph(f"Processing Summary:")
    summary.runs[0].bold = True
    doc.add_paragraph(f"• Images processed: {len(new_images)}")
    doc.add_paragraph(f"• Total words: {len(combined_text.split())}")
    avg_confidence = total_confidence / len(new_images)
    doc.add_paragraph(f"• Average confidence: {avg_confidence:.0%}")
    doc.add_paragraph(f"• Processed: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # Save document
    output_path = os.path.join(OUTPUT_FOLDER, filename)
    doc.save(output_path)

    log_message(f"Created: {filename}")
    log_message(f"Total words extracted: {len(combined_text.split())}")
    log_message(f"Average confidence: {avg_confidence:.0%}")

if __name__ == "__main__":
    try:
        process_images()
    except Exception as e:
        log_message(f"ERROR: {str(e)}")
        raise
